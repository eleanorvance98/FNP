"""
Official Hybrid option: CSE discovery + selective download + Google Cloud translation + CSV & DOCX export

What it does:
- Uses Google Custom Search JSON API (Programmable Search Engine) to discover URLs using Sinhala queries
  plus Google date operators in the query string (after:/before:) for 2015–2022.
- Normalizes URLs (including converting translate.goog proxy URLs to canonical lankadeepa.lk URLs).
- Downloads and parses articles from canonical URLs.
- Translates Sinhala -> English using Google Cloud Translation API (v2 client).
- Selectively KEEP only those articles whose English translation strongly matches your topic keywords.
- Writes:
  * master CSV (one row/article; includes Sinhala + English text)
  * DOCX batches (e.g., 500 articles per doc)

Notes:
- Google's Custom Search API supports dateRestrict, but it's relative (d/w/m/y), not a fixed year span. 
  So we put after:/before: directly in the query string (Google query operators).
- The API has a free tier (commonly 100 queries/day), and supports pagination via start=1,11,21,...
- Your selective download step keeps the actual site load manageable.

You MUST review and comply with any site Terms of Use. Robots.txt is not the only constraint.
"""

from __future__ import annotations

import csv
import os
from pydoc import doc
import re
import json
import time
import html
import hashlib
from datetime import datetime
from typing import Dict, List, Tuple

import requests
import pandas as pd
from bs4 import BeautifulSoup
from tqdm import tqdm
from docx import Document

from google.cloud import translate_v2 as translate  # google-cloud-translate==2.0.1

def clean_for_csv(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    # Remove nulls that can break Excel
    s = s.replace("\x00", "")
    # Keep paragraphs but avoid literal newlines that confuse some CSV readers
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    return s

# -----------------------------
# Configuration
# -----------------------------
CSE_API_KEY = os.environ.get("CSE_API_KEY", "").strip() #environmental variables for API  keys
CSE_CX = os.environ.get("CSE_CX", "").strip()

OUT_DIR = "lankadeepa_hybrid_output"
MASTER_CSV = os.path.join(OUT_DIR, "master_lankadeepa_2022_2022.csv")
MASTER_XLSX = os.path.join(OUT_DIR, "master_lankadeepa_2022_2022.xlsx")
STATE_JSON = os.path.join(OUT_DIR, "state.json")

DOCX_BATCH_SIZE = 500

# Polite delay between *article* fetches (not between CSE calls).
# Since we're downloading selectively, a modest delay is usually fine.
# If you want to be ultra-conservative, increase this.
ARTICLE_FETCH_DELAY_SECONDS = 5

# Polite delay between CSE calls (to avoid hammering Google API & to be nice)
CSE_DELAY_SECONDS = 0.5

# Identify yourself
HEADERS = {
    "User-Agent": "AcademicResearchBot/1.0 (contact: your_email@example.com)"
}

# -----------------------------
# Topic matching (English, after translation)
# -----------------------------
TOPIC_KEYWORDS_EN = [
    "farmer protest",
    "farmers protest",
    "farmers' protest",
    "farmers protests",
    "agricultural protest",
    "protest by farmers",
    "farmers demonstration",
    "farmers rally",
    "fertilizer ban",
    "import ban",
    "ban protest"
]

# You can tune these thresholds:
MIN_TOTAL_KEYWORD_HITS = 1          # minimum count across all keywords CHANGE TO 2 AFTER TEST
MIN_DISTINCT_KEYWORDS = 2           # minimum number of distinct keyword phrases present #CHANGE TO 1 IF NEED MORE HITS
MIN_ENGLISH_TEXT_CHARS = 300        # ignore very short pages / non-articles    CHANGE TO 800 AFTER TEST

# -----------------------------
# CSE discovery settings
# -----------------------------
DOMAIN = "lankadeepa.lk"

# For each year, we run a fixed number of pages (10 results per page).
# One "page" here = one API call.
MAX_PAGES_PER_QUERY = 1  # 10 pages -> up to 100 URLs per query. SET TO 1 FOR TEST, CHANGE TO 10 AFTER CONFIRMING IT WORKS!!

# If you want broader recall, provide multiple Sinhala query variants.
# Replace the Sinhala terms below with the exact ones you used.
SINHALA_QUERY_VARIANTS = [
    # Example placeholders — replace with your actual Sinhala terms for farmer + protest:
    '"ගොවීන්" "උද්ඝෝෂණ"',       # farmers + protest
    '"ගොවි" "විරෝධතාව"',        # farmer + demonstration/protest
    '"ගොවි" "උද්ඝෝෂණ"',          # farmer + protest
]

START_YEAR = 2022 #change to 2020 after test run
END_YEAR = 2022  # inclusive


# -----------------------------
# Utilities
# -----------------------------
def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

def sha_id(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8")).hexdigest()[:12]

def load_state() -> Dict:
    if os.path.exists(STATE_JSON):
        with open(STATE_JSON, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "discovered_urls": {},  # url -> metadata
        "processed_urls": {},   # url -> status + outputs
    }

def save_state(state: Dict) -> None:
    with open(STATE_JSON, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

def init_master_csv() -> pd.DataFrame:
    if os.path.exists(MASTER_CSV):
        return pd.read_csv(MASTER_CSV)
    return pd.DataFrame(columns=[
        "url",
        "url_canonical",
        "year",
        "query_variant",
        "cse_title",
        "cse_snippet",
        "title",
        "date_raw",
        "body_si",
        "body_en",
        "keyword_hits_total",
        "keyword_hits_distinct",
        "docx_file",
        "docx_article_index"
    ])

def append_to_master_csv(df: pd.DataFrame, row: Dict) -> pd.DataFrame:
    # Clean long text fields before writing
    row["body_si"] = clean_for_csv(row.get("body_si", ""))
    row["body_en"] = clean_for_csv(row.get("body_en", ""))

    df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)

    df.to_csv(
        MASTER_CSV,
        index=False,
        encoding="utf-8-sig",          # BOM so Excel reads Sinhala correctly
        quoting=csv.QUOTE_ALL,         # safest for commas/newlines
        escapechar="\\",               # extra safety
        lineterminator="\n"
    )
    return df


# -----------------------------
# URL normalization
# -----------------------------
def canonicalize_lankadeepa_url(url: str) -> str:
    """
    Converts:
      - translate.goog proxy URLs -> canonical lankadeepa.lk URLs
      - strips Google Translate query params (_x_tr_*)
      - normalizes scheme/host
    """
    u = url.strip()

    # If it's a translate.goog host like: https://www-lankadeepa-lk.translate.goog/news/...?..._x_tr...
    m = re.match(r"^https?://www-lankadeepa-lk\.translate\.goog(?P<path>/.*)$", u)
    if m:
        path_and_q = m.group("path")
        # Remove translate params (everything after ?)
        path = path_and_q.split("?", 1)[0]
        return f"https://www.lankadeepa.lk{path}"

    # If it's already lankadeepa, strip any translate params
    if "lankadeepa.lk" in u:
        base = u.split("?", 1)[0]
        # Ensure canonical host
        base = re.sub(r"^https?://(www\.)?lankadeepa\.lk", "https://www.lankadeepa.lk", base)
        return base

    return u

# -----------------------------
# CSE (Custom Search JSON API)
# -----------------------------
def cse_list(q: str, start: int = 1) -> Dict:
    """
    Calls Custom Search JSON API.
    """
    if not CSE_API_KEY or not CSE_CX:
        raise RuntimeError("Missing CSE_API_KEY or CSE_CX. Set them as environment variables.")

    endpoint = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": CSE_API_KEY,
        "cx": CSE_CX,
        "q": q,
        "start": start,
        "num": 10,
    }
    r = requests.get(endpoint, params=params, timeout=30)
    r.raise_for_status()
    return r.json()

def build_year_queries() -> List[Tuple[int, str, str]]:
    """
    Returns list of (year, query_variant, query_string)
    Uses Google query operators after:/before: to approximate a fixed window.
    """
    queries = []
    for year in range(START_YEAR, END_YEAR + 1):
        after = f"{year}-01-01"
        before = f"{year+1}-01-01"
        for variant in SINHALA_QUERY_VARIANTS:
            q = f"site:{DOMAIN} {variant} after:{after} before:{before}"
            queries.append((year, variant, q))
    return queries

def discover_urls_via_cse(state: Dict) -> int:
    """
    Discovers URLs and saves them into state["discovered_urls"].
    Returns number of NEW canonical URLs added.
    """
    new_count = 0
    queries = build_year_queries()

    for year, variant, q in tqdm(queries, desc="CSE discovery"):
        for page in range(MAX_PAGES_PER_QUERY):
            start = 1 + page * 10
            data = cse_list(q, start=start)
            items = data.get("items", [])
            if not items:
                break

            for it in items:
                link = it.get("link", "")
                title = it.get("title", "")
                snippet = it.get("snippet", "")

                if not link:
                    continue

                canon = canonicalize_lankadeepa_url(link)

                # Keep only domain pages (canonical must contain lankadeepa.lk)
                if "lankadeepa.lk" not in canon:
                    continue

                if canon not in state["discovered_urls"]:
                    state["discovered_urls"][canon] = {
                        "year": year,
                        "query_variant": variant,
                        "cse_title": title,
                        "cse_snippet": snippet,
                        "original_link": link,
                    }
                    new_count += 1

            save_state(state)
            time.sleep(CSE_DELAY_SECONDS)

    return new_count

# -----------------------------
# Article fetching & parsing
# -----------------------------
def fetch_html(url: str) -> str:
    r = requests.get(url, headers=HEADERS, timeout=30)
    r.raise_for_status()
    return r.text

def extract_article_fields(html_text: str) -> Tuple[str, str, str]:
    """
    Extract (title, date_raw, body_text_sinhala) from HTML.

    IMPORTANT:
    - This uses generic selectors and may need tuning for Lankadeepa.
    - Start by testing on 5–10 pages and adjust the container selectors to reduce noise.
    """
    soup = BeautifulSoup(html_text, "lxml")

    # Title
    h1 = soup.find("h1")
    title = h1.get_text(" ", strip=True) if h1 else ""
    if not title and soup.title:
        title = soup.title.get_text(" ", strip=True)

    # Date attempts
    date_raw = ""
    time_tag = soup.find("time")
    if time_tag:
        date_raw = time_tag.get_text(" ", strip=True)

    if not date_raw:
        meta_selectors = [
            ('meta', {'property': 'article:published_time'}),
            ('meta', {'name': 'publish-date'}),
            ('meta', {'name': 'pubdate'}),
            ('meta', {'name': 'date'}),
            ('meta', {'property': 'og:updated_time'}),
        ]
        for tag, attrs in meta_selectors:
            m = soup.find(tag, attrs=attrs)
            if m and m.get("content"):
                date_raw = m["content"].strip()
                break

    if not date_raw:
        for sel in [".date", ".post-date", ".publish-date", ".article-date", ".news-date", ".timestamp"]:
            el = soup.select_one(sel)
            if el and el.get_text(strip=True):
                date_raw = el.get_text(" ", strip=True)
                break

    # Body container attempts
    body_text = ""
    container_selectors = [
        "article",
        ".article-body",
        ".post-content",
        ".content",
        ".details-content",
        ".news-content",
        ".entry-content",
    ]
    for sel in container_selectors:
        container = soup.select_one(sel)
        if not container:
            continue
        paragraphs = [p.get_text(" ", strip=True) for p in container.find_all("p")]
        paragraphs = [p for p in paragraphs if p and len(p) > 20]
        if len(paragraphs) >= 3:
            body_text = "\n".join(paragraphs)
            break

    # Fallback: pick many <p> (noisier)
    if not body_text:
        paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
        paragraphs = [p for p in paragraphs if p and len(p) > 20]
        body_text = "\n".join(paragraphs[:80])

    # Unescape any HTML entities
    title = html.unescape(title)
    date_raw = html.unescape(date_raw)
    body_text = html.unescape(body_text)

    return title, date_raw, body_text

# -----------------------------
# Translation (Google Cloud)
# -----------------------------
def chunk_text(text: str, max_chars: int = 4500) -> List[str]:
    text = text.strip()
    if len(text) <= max_chars:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        # try splitting on newline for cleaner boundaries
        nl = text.rfind("\n", start, end)
        if nl > start + 500:
            end = nl
        chunks.append(text[start:end].strip())
        start = end
    return [c for c in chunks if c]

def translate_si_to_en(client: translate.Client, text: str) -> str:
    if not text.strip():
        return ""
    parts = chunk_text(text)
    out = []
    for p in parts:
        resp = client.translate(p, target_language="en")  # auto-detect source
        translated = resp.get("translatedText", "")
        # IMPORTANT: Google often returns HTML entities like &#39;
        translated = html.unescape(translated)  # <-- unescape entities
        out.append(translated)
    return "\n".join(out)

# -----------------------------
# Topic matching (strong match)
# -----------------------------
def count_keyword_hits(text_en: str) -> Tuple[int, int]:
    """
    Returns (total_hits, distinct_keywords_found).
    Simple substring counting; you can swap to regex/lemmatization if desired.
    """
    t = (text_en or "").lower()
    total = 0
    distinct = 0
    for kw in TOPIC_KEYWORDS_EN:
        hits = t.count(kw.lower())
        if hits > 0:
            distinct += 1
            total += hits
    return total, distinct

def is_strong_match(text_en: str) -> Tuple[bool, int, int]:
    if not text_en or len(text_en) < MIN_ENGLISH_TEXT_CHARS:
        return False, 0, 0
    total, distinct = count_keyword_hits(text_en)
    ok = (total >= MIN_TOTAL_KEYWORD_HITS) and (distinct >= MIN_DISTINCT_KEYWORDS)
    return ok, total, distinct

# -----------------------------
# DOCX batching
# -----------------------------
def docx_path_for_batch(batch_num: int) -> str:
    return os.path.join(OUT_DIR, f"lankadeepa_2022_2022_part{batch_num:03d}.docx") #Change file name after test run to reflect real time period


def load_or_create_doc(doc_path: str) -> Document:
    if os.path.exists(doc_path):
        return Document(doc_path)
    doc = Document()
    doc.add_heading("Lankadeepa: Farmer-protest-related articles (2022–2022)", level=1)
    return doc

# -----------------------------
# Main selective downloader
# -----------------------------
def process_and_export(state: Dict) -> None:
    ensure_dir(OUT_DIR)

    df = init_master_csv()
    translate_client = translate.Client()

    # Determine current docx batch position from CSV rows already saved
    kept_so_far = df["docx_file"].notna().sum()
    batch_num = kept_so_far // DOCX_BATCH_SIZE + 1
    in_batch_index = kept_so_far % DOCX_BATCH_SIZE

    doc_path = docx_path_for_batch(batch_num)
    doc = load_or_create_doc(doc_path)

    urls = list(state["discovered_urls"].keys())
    print(f"Discovered URLs in state: {len(urls)}")

    for canon_url in tqdm(urls, desc="Selective download"):
        if canon_url in state["processed_urls"]:
            continue

        meta = state["discovered_urls"][canon_url]
        year = meta.get("year", "")
        variant = meta.get("query_variant", "")
        cse_title = meta.get("cse_title", "")
        cse_snippet = meta.get("cse_snippet", "")
        original_link = meta.get("original_link", "")

        try:
            html_text = fetch_html(canon_url)
            title, date_raw, body_si = extract_article_fields(html_text)

            # Translate to English (Google Cloud)
            body_en = translate_si_to_en(translate_client, body_si)

            # Strong match test
            keep, hits_total, hits_distinct = is_strong_match(body_en)

            # Record to CSV always (so you can audit); only write to DOCX if keep==True
            row = {
                "url": original_link or canon_url,
                "url_canonical": canon_url,
                "year": year,
                "query_variant": variant,
                "cse_title": cse_title,
                "cse_snippet": cse_snippet,
                "title": title,
                "date_raw": date_raw,
                "body_si": body_si,
                "body_en": body_en,
                "keyword_hits_total": hits_total,
                "keyword_hits_distinct": hits_distinct,
                "docx_file": None,
                "docx_article_index": None,
            }

            if keep:
                in_batch_index += 1
                row["docx_file"] = os.path.basename(doc_path)
                row["docx_article_index"] = in_batch_index

                # Append to DOCX (English text for analysis consistency)
                doc.add_heading(title or "(no title)", level=2)
                doc.add_paragraph(f"Date: {date_raw}")
                doc.add_paragraph(f"URL: {canon_url}")
                doc.add_paragraph("")
                doc.add_paragraph(body_en)
                doc.add_page_break()

                # Rotate DOCX
                if in_batch_index >= DOCX_BATCH_SIZE:
                    doc.save(doc_path)
                    batch_num += 1
                    in_batch_index = 0
                    doc_path = docx_path_for_batch(batch_num)
                    doc = load_or_create_doc(doc_path)

            df = append_to_master_csv(df, row)

            state["processed_urls"][canon_url] = {
                "kept": keep,
                "keyword_hits_total": hits_total,
                "keyword_hits_distinct": hits_distinct,
                "docx_file": row["docx_file"],
                "docx_article_index": row["docx_article_index"],
                "timestamp_processed_utc": datetime.utcnow().isoformat(timespec="seconds") + "Z",
            }
            save_state(state)

            time.sleep(ARTICLE_FETCH_DELAY_SECONDS)

        except Exception as e:
            # Record error and continue
            state["processed_urls"][canon_url] = {
                "kept": False,
                "error": str(e),
                "timestamp_processed_utc": datetime.utcnow().isoformat(timespec="seconds") + "Z",
            }
            save_state(state)
            continue

    # ---- END OF URL LOOP ----
    doc.save(doc_path)
    df.to_excel(MASTER_XLSX, index=False)

    print("Done. Outputs:")
    print(" -", MASTER_CSV)
    print(" -", MASTER_XLSX)
    print(" - DOCX files in", OUT_DIR)


def main():
    ensure_dir(OUT_DIR)
    state = load_state()

    print("Starting discovery via CSE…")
    new_urls = discover_urls_via_cse(state)
    print(f"Discovery complete. New canonical URLs added: {new_urls}")

    print("Starting selective download + translate + export…")
    process_and_export(state)


if __name__ == "__main__":
    main()

        


   


